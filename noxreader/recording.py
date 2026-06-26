"""
noxreader.recording
核心讀取實作：使用 numpy.memmap 達成最佳化開檔與區段讀取。
支援 Nox PSG 儀器產生的 raw .ndf + SETUP.INI 為主的資料格式。
"""

from __future__ import annotations
import os
import re
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Union
import numpy as np
import sqlite3
import configparser
from datetime import datetime, timedelta

from .nox_edf import NoxEdfHeader, parse_nox_edf, read_nox_edf_channel, try_pyedflib_open

PathLike = Union[str, Path]


def format_channel_source(entry: Dict) -> str:
    """回傳通道資料來源標籤，例如 NDF_c3.ndf 或 EDF_xxx.edf。"""
    path = Path(str(entry.get("path", "")))
    fname = path.name if path.name else str(entry.get("path", ""))
    if entry.get("source") == "edf" or path.suffix.lower() == ".edf":
        return f"EDF_{fname}"
    return f"NDF_{fname}"


def _normalize(name: str) -> str:
    """將通道名稱正規化，用於比對檔名與 SETUP 描述。"""
    n = name.lower().strip()
    n = re.sub(r"[\s_\-]+", "", n)  # 移除空白、底線、連字號
    n = n.replace("impedance", "imp")
    n = n.replace("ambientlight", "light")
    n = n.replace("audio", "audio")
    return n


def _guess_fs_unit_from_label(label: str) -> tuple[float, str]:
    """從 label 猜 fs 和 unit（用於 lenient fallback 當 pyedflib 無法開啟壞 EDF 時）。"""
    l = (label or '').lower()
    if 'imped' in l or 'imp' in l:
        return 1.0, 'Ohm'
    if any(x in l for x in ['c3', 'c4', 'f3', 'f4', 'o1', 'o2', 'e1', 'e2', 'm1', 'm2', 'ecg', 'eeg']):
        return 200.0, 'V'
    if any(x in l for x in ['rip', 'abdomen', 'thora', 'chest', 'flow', 'cflow', 'induct']):
        return 25.0, 'V'
    if 'spo2' in l or 'saturation' in l:
        return 3.0, '%'
    if 'pulse' in l:
        return 75.0, ''
    if 'light' in l:
        return 1.0, 'lx'
    if 'pos' in l or 'angle' in l:
        return 20.0, 'deg'
    if 'voltage' in l:
        return 100.0, 'V'
    if 'heart' in l or 'rate' in l:
        return 1.0, 'bpm'
    if 'snore' in l:
        return 200.0, 'V'
    if 'mask' in l or 'pressure' in l:
        return 25.0, 'cmH2O'
    return 25.0, 'V'


_EVENT_XLS_GLOBS = (
    '*event*.xls', '*Event*.xls', '*EVENT*.xls',
    '*event*.xlsx', '*Event*.xlsx', '*EVENT*.xlsx',
)
_XLSX_NS = {"x": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
_MIN_RELIABLE_NDF_DURATION_SEC = 60.0

# Nox .ndf 樣本格式 <Format> -> (numpy dtype, bytes/sample)。
# 實測四位病患（D02/D03/D18/D20）僅出現這四種；非 int16 的通道在舊版被當 int16 讀，
# 導致波形數值錯誤、duration 倍增（int32→2x、byte→0.5x）。預設退回 int16。
_NDF_FORMAT_DTYPE: Dict[str, Tuple[str, int]] = {
    "int16": ("<i2", 2),
    "uint16": ("<u2", 2),
    "int32": ("<i4", 4),
    "uint32": ("<u4", 4),
    "byte": ("<u1", 1),
    "uint8": ("<u1", 1),
    "sbyte": ("<i1", 1),
    "int8": ("<i1", 1),
    "single": ("<f4", 4),
    "float": ("<f4", 4),
    "double": ("<f8", 8),
}
_DEFAULT_NDF_DTYPE = ("<i2", 2)

# 標記 Nox 資料區起點的 tag：02 02 08 00 00 00 後接 float64 取樣率，再接 01 02 + 4-byte len，之後即樣本資料。
_NDF_DATA_ANCHOR = bytes.fromhex("020208000000")
_NDF_ANCHOR_TO_DATA = 6 + 8 + 2 + 4  # anchor(6) + float64 fs(8) + 01 02(2) + len(4) = 20 bytes


def _read_standard_edf_meta(edf_path: PathLike) -> Tuple[float, Optional[datetime]]:
    """從標準 EDF 讀取總時長與開始時間（pyedflib）。"""
    p = Path(edf_path)
    if not p.exists():
        return 0.0, None
    try:
        import pyedflib

        f = pyedflib.EdfReader(str(p))
        dur = float(f.file_duration or 0.0)
        start = f.getStartdatetime()
        f.close()
        return dur, start
    except Exception:
        return 0.0, None


def _edf_duration_sec(edf_path: PathLike) -> float:
    """取得 EDF 檔案時長：先試 pyedflib，再試 Nox parser。"""
    dur, _ = _read_standard_edf_meta(edf_path)
    if dur > 0:
        return dur
    try:
        return float(parse_nox_edf(edf_path).duration_sec)
    except Exception:
        return 0.0


def _read_xlsx_rows_via_xml(xlsx_path: Path) -> List[List]:
    """直接解析 xlsx 內 sheet XML（相容 Nox 匯出檔的非法儲存格座標 r=\"num\"）。"""
    with zipfile.ZipFile(xlsx_path) as zf:
        sheet_names = [
            n for n in zf.namelist()
            if n.startswith("xl/worksheets/") and n.endswith(".xml")
        ]
        if not sheet_names:
            return []
        root = ET.fromstring(zf.read(sheet_names[0]))
    rows: List[List] = []
    for row_el in root.findall(".//x:sheetData/x:row", _XLSX_NS):
        vals: List = []
        for cell in row_el.findall("x:c", _XLSX_NS):
            v_el = cell.find("x:v", _XLSX_NS)
            if v_el is None or not (v_el.text or "").strip():
                vals.append(None)
                continue
            text = v_el.text.strip()
            if cell.get("t") == "str":
                vals.append(text)
            else:
                try:
                    vals.append(float(text))
                except ValueError:
                    vals.append(text)
        if vals:
            rows.append(vals)
    return rows


def _load_xlsx_event_table(xlsx_path: Path) -> List[Tuple]:
    """讀取 xlsx 事件表（第 3 列起），openpyxl 失敗時退回 XML。"""
    rows_out: List[Tuple] = []
    try:
        import openpyxl

        wb = openpyxl.load_workbook(str(xlsx_path), data_only=True)
        sheet = wb.active
        for row in sheet.iter_rows(min_row=3, values_only=True):
            if not row:
                continue
            padded = list(row) + [None] * max(0, 4 - len(row))
            rows_out.append(tuple(padded[:4]))
        wb.close()
        return rows_out
    except Exception as exc:
        print(f"[XLS] openpyxl 無法讀取 {xlsx_path.name} ({exc})，改以 XML 解析")

    raw_rows = _read_xlsx_rows_via_xml(xlsx_path)
    if len(raw_rows) < 3:
        return []
    for row in raw_rows[2:]:
        padded = list(row) + [None] * max(0, 4 - len(row))
        rows_out.append(tuple(padded[:4]))
    return rows_out


def _open_xlrd_workbook_tolerant(xls_path: Path):
    """開啟 Nox 匯出的 .xls；略過損壞 WriteAccess 造成的 utf-16 解碼錯誤。"""
    import xlrd
    import xlrd.book

    if not getattr(_open_xlrd_workbook_tolerant, "_patched", False):
        _orig = xlrd.book.Book.handle_writeaccess

        def _tolerant_writeaccess(self, data):
            try:
                _orig(self, data)
            except UnicodeDecodeError:
                pass

        xlrd.book.Book.handle_writeaccess = _tolerant_writeaccess
        _open_xlrd_workbook_tolerant._patched = True

    _devnull = open(os.devnull, "w")
    try:
        return xlrd.open_workbook(str(xls_path), logfile=_devnull, on_demand=True)
    finally:
        try:
            _devnull.close()
        except Exception:
            pass


def _load_xls_event_table(xls_path: Path) -> List[Tuple]:
    """讀取 .xls 事件表（第 3 列起）。"""
    wb = _open_xlrd_workbook_tolerant(xls_path)
    sheet = wb.sheet_by_index(0)
    if sheet.nrows < 3:
        return []
    rows_out: List[Tuple] = []
    for r in range(2, sheet.nrows):
        padded = [
            sheet.cell(r, c).value if c < sheet.ncols else None
            for c in range(4)
        ]
        rows_out.append(tuple(padded))
    return rows_out


def _events_from_table_rows(
    table_rows: List[Tuple],
    rec_start: datetime,
    *,
    id_prefix: str = "xls",
) -> List[Dict]:
    """將事件表列轉成 parsed event dict 列表。"""
    excel_epoch = datetime(1899, 12, 30)
    events: List[Dict] = []
    for r_idx, row in enumerate(table_rows, start=2):
        try:
            event_name = str(row[0] or "").strip()
            if not event_name or event_name.lower().startswith("analysis"):
                continue
            dur = float(row[1] or 0)
            start_val = row[2]
            end_val = row[3]
            if isinstance(start_val, datetime):
                start_dt = start_val
            elif isinstance(start_val, (int, float)):
                start_dt = excel_epoch + timedelta(days=float(start_val))
            else:
                continue
            if isinstance(end_val, datetime):
                end_dt = end_val
            elif isinstance(end_val, (int, float)):
                end_dt = excel_epoch + timedelta(days=float(end_val))
            else:
                end_dt = start_dt
            rel_start = (start_dt - rec_start).total_seconds()
            rel_end = None
            if end_dt and (end_dt - start_dt).total_seconds() > 0.0005:
                computed = (end_dt - rec_start).total_seconds()
                if computed > rel_start + 0.0005:
                    rel_end = computed
            if rel_start < 0:
                rel_start = 0
            events.append({
                "id": f"{id_prefix}_{r_idx}",
                "starts_at": 0,
                "ends_at": 0,
                "type": event_name,
                "location": event_name,
                "notes": f"來自 XLS 分析 (dur={dur:.2f}s)",
                "rel_start": rel_start,
                "rel_end": rel_end,
                "_source": "xls",
            })
        except Exception:
            continue
    return events


def find_event_xls_paths(
    patient_dir: Optional[PathLike],
    raw_dir: Optional[PathLike] = None,
) -> List[Path]:
    """搜尋 Nox 分析軟體輸出的 event Excel，優先 patient_dir 且檔名含 event、較大檔案。"""
    seen: set[Path] = set()
    cands: List[Path] = []
    patient_p = Path(patient_dir).resolve() if patient_dir else None
    for base in (patient_dir, raw_dir):
        if not base:
            continue
        bp = Path(base)
        if not bp.exists():
            continue
        for pat in _EVENT_XLS_GLOBS:
            for p in bp.glob(pat):
                if p.is_file():
                    rp = p.resolve()
                    if rp not in seen:
                        seen.add(rp)
                        cands.append(rp)
    def _score(p: Path) -> tuple:
        in_event = 'event' in p.name.lower()
        in_patient = bool(patient_p and p.parent.resolve() == patient_p)
        try:
            size = p.stat().st_size
        except OSError:
            size = 0
        return (in_event, in_patient, size)
    cands.sort(key=_score, reverse=True)
    return cands


def parse_event_xls_file(
    xls_path: PathLike,
    rec_start: Optional[datetime],
    *,
    fallback_start: Optional[datetime] = None,
) -> List[Dict]:
    """解析單一 event_*.xls / *.xlsx，回傳含 rel_start/rel_end 的事件列表。"""
    xls_path = Path(xls_path)
    if not xls_path.exists():
        return []
    if rec_start is None:
        rec_start = fallback_start
    if rec_start is None:
        # start_datetime 現已可靠解析（見 nox_edf._parse_start_datetime），正常不會走到這裡。
        # 沒有任何起始時間就無法計算事件的相對位置，誠實略過而非用假的硬編碼日期。
        print(f"[XLS] 缺少錄音起始時間，無法計算事件相對位置，略過：{xls_path.name}")
        return []
    events: List[Dict] = []
    ext = xls_path.suffix.lower()
    try:
        if ext in ('.xlsx', '.xlsm'):
            table_rows = _load_xlsx_event_table(xls_path)
        else:
            try:
                import xlrd  # noqa: F401
            except ImportError:
                print(
                    f"[XLS] xlrd 未安裝，無法讀取舊版 .xls 事件檔：{xls_path.name}。"
                    "請執行: pip install xlrd==1.2.0"
                )
                return []
            table_rows = _load_xls_event_table(xls_path)
        events = _events_from_table_rows(table_rows, rec_start)
        if events:
            print(f"[XLS] 已解析 {len(events)} 筆事件：{xls_path.name}")
        return events
    except Exception as e:
        print(f"[XLS] 解析失敗 {xls_path}: {e}")
        return []


def find_raw_data_dir(patient_dir: PathLike) -> Optional[Path]:
    """在病患資料夾中找到包含 SETUP.INI 與大量 .ndf 的 raw data 子目錄。"""
    p = Path(patient_dir)
    candidates = []
    # 常見結構: patient/raw data_xxx/timestamp - hash/
    for sub in p.rglob("*"):
        if sub.is_dir():
            setup = sub / "SETUP.INI"
            if setup.exists():
                ndf_count = len(list(sub.glob("*.ndf")))
                if ndf_count > 5:
                    candidates.append((ndf_count, sub))
    if not candidates:
        return None
    # 選 ndf 最多的
    candidates.sort(reverse=True)
    return candidates[0][1]


def parse_setup_ini(setup_path: PathLike) -> Dict[str, Dict]:
    """
    解析 SETUP.INI。
    回傳： {internal_key: {'display': , 'desc':, 'fs': float, 'unit': str, 'raw_line': }}
    例如 EXG8 -> C3 / EEG-C3 / 200 / V
    """
    setup_path = Path(setup_path)
    meta: Dict[str, Dict] = {}
    if not setup_path.exists():
        return meta
    with open(setup_path, "r", encoding="utf-8", errors="ignore") as f:
        for line in f:
            line = line.strip()
            if not line or "=" not in line:
                continue
            key, val = line.split("=", 1)
            key = key.strip()
            val = val.strip()
            if not val:
                continue
            parts = [x.strip() for x in val.split(";")]
            display = parts[0] if parts else key
            desc = parts[1] if len(parts) > 1 else ""
            try:
                fs = float(parts[2]) if len(parts) > 2 and parts[2] else 0.0
            except ValueError:
                fs = 0.0
            unit = parts[3] if len(parts) > 3 else ""
            meta[key] = {
                "display": display,
                "desc": desc,
                "fs": fs,
                "unit": unit,
                "raw": val,
            }
    return meta


def build_channel_index(raw_dir: Path, setup_meta: Dict[str, Dict]) -> Dict[str, Dict]:
    """
    掃描 raw_dir 內所有 .ndf，建立可存取的 channel 索引。
    優先使用 SETUP 裡的 display 名稱對應，否則退回使用檔名 stem。
    每個 entry 包含: path, fs, unit, desc, samples, duration_sec, internal_key
    """
    index: Dict[str, Dict] = {}
    ndf_files = sorted(raw_dir.glob("*.ndf"))
    # 建立反向查詢：正規化後的 display 名稱 -> meta
    display_map: Dict[str, Dict] = {}
    for int_key, m in setup_meta.items():
        disp = m["display"]
        nk = _normalize(disp)
        display_map[nk] = {**m, "internal_key": int_key}
        # 也支援簡短如 c3, spo2（spo2 類通道現在優先由 ndf 內 header 提供正確 fs=3 全長）
        short = _normalize(disp.split()[0] if " " in disp else disp)
        if short and short not in display_map:
            display_map[short] = {**m, "internal_key": int_key}

    for ndf in ndf_files:
        stem = ndf.stem  # e.g. "c3", "abdomen rip", "ecg"
        nk = _normalize(stem)
        entry = {
            "path": ndf,
            "stem": stem,
            "fs": 0.0,
            "unit": "",
            "desc": stem,
            "internal_key": None,
            "samples": 0,
            "duration_sec": 0.0,
            "header_offset": 0,
            "inferred_fs": None,
            "np_dtype": _DEFAULT_NDF_DTYPE[0],
            "bytes_per_sample": _DEFAULT_NDF_DTYPE[1],
            "data_start_byte": None,
            "start_datetime": None,   # 該通道第一筆 sample 的絕對時間（嵌入 header）
            "start_offset_sec": 0.0,  # 相對全域原點的偏移（B2 渲染用），由 _compute_time_origin 填
        }

        # 優先從 .ndf 自身 embedded header 解析（最準確，尤其 Nonin oximeter 的 spo2 / spo2 b-b / pleth / pulse）
        # 可避免 SETUP partial match 錯誤（例如 "2"  substring 誤配進 "spo2" 拿到 200Hz）
        ndf_meta = _parse_ndf_channel_meta(ndf)
        # dtype / 資料起點不依賴 fs（int32/byte 通道即使 fs 來自 setup 也須正確套用）
        if ndf_meta.get("np_dtype"):
            entry["np_dtype"] = ndf_meta["np_dtype"]
            entry["bytes_per_sample"] = ndf_meta.get("bytes_per_sample", 2)
        if ndf_meta.get("format"):
            entry["format"] = ndf_meta["format"]
        if ndf_meta.get("data_start_byte") is not None:
            entry["data_start_byte"] = ndf_meta["data_start_byte"]
        if ndf_meta.get("start_datetime") is not None:
            entry["start_datetime"] = ndf_meta["start_datetime"]
        if ndf_meta.get("fs") is not None:
            entry["fs"] = ndf_meta["fs"]
            if ndf_meta.get("unit"):
                entry["unit"] = ndf_meta["unit"]
            if ndf_meta.get("label"):
                entry["desc"] = ndf_meta["label"]
            if ndf_meta.get("data_start_hint"):
                entry["_data_start_hint"] = ndf_meta["data_start_hint"]
            if ndf_meta.get("scale") is not None:
                entry["scale"] = ndf_meta["scale"]
            if ndf_meta.get("type"):
                entry["type"] = ndf_meta["type"]

        # 嘗試從 setup 取得精確資訊（作為補充，特別是 display name 對應 EXGxx）
        if entry["fs"] <= 0 and nk in display_map:
            m = display_map[nk]
            entry.update(
                {
                    "fs": m["fs"],
                    "unit": m["unit"] or entry.get("unit", ""),
                    "desc": m.get("desc", stem) or entry.get("desc", stem),
                    "internal_key": m.get("internal_key"),
                }
            )
        elif entry["fs"] <= 0:
            # 嚴格化部分比對：排除極短 k（如 "2","1"）或純數字，避免 "2" in "spo2" 之類誤配
            for k, m in display_map.items():
                if (k in nk or nk in k) and len(k) >= 3 and not k.isdigit():
                    entry.update(
                        {
                            "fs": m["fs"],
                            "unit": m["unit"] or entry.get("unit", ""),
                            "desc": m.get("desc", stem) or entry.get("desc", stem),
                            "internal_key": m.get("internal_key"),
                        }
                    )
                    break

        # 從檔案大小推算 samples 與 duration（即使 fs 未知也先記錄）
        # 結構化路徑：samples = (fsize - data_start_byte) // bytes_per_sample，已精確扣除 header，無需 _apply_header_offsets。
        # legacy 路徑（無 anchor，如合成測試檔）：沿用 fsize//bps（預設 int16），header 校正交給 _apply_header_offsets。
        try:
            fsize = ndf.stat().st_size
            bps = entry.get("bytes_per_sample", 2) or 2
            ds = entry.get("data_start_byte")
            if ds is not None:
                samples = max(0, (fsize - ds) // bps)
            else:
                samples = fsize // bps
            entry["samples"] = samples
            if entry["fs"] > 0:
                entry["duration_sec"] = samples / entry["fs"]
        except Exception:
            pass

        # 以 stem 與正規化後名稱同時註冊，方便使用者用不同方式存取
        index[stem] = entry
        if nk and nk not in index:
            index[nk] = entry
        # 也用 desc 的一部分
        if entry["desc"] and entry["desc"] not in index:
            index[entry["desc"]] = entry

    return index


def compute_overall_duration(ch_index: Dict[str, Dict]) -> float:
    """取各通道推算 duration 的中位數作為 recording 總時長。"""
    durs = [e["duration_sec"] for e in ch_index.values() if e.get("duration_sec", 0) > 60]
    if not durs:
        return 0.0
    durs.sort()
    return durs[len(durs) // 2]


def _detect_header_offset(path: PathLike) -> int:
    """
    偵測 .ndf 檔案開頭的 XML junk header 偏移量（以 int16 samples 計）。
    掃描低 std + 有界值（避開開頭高值 junk 如 20302），適用 position 等低頻道。
    增強版：若原條件（高 std）在低變異通道（如 spo2 @3Hz，值穩定在 900 左右，std 小）抓不到，
    則退而尋找第一個「合理訊號區」（連續 32+ 筆 |v|<2000 的起始），確保 spo2 / pleth 等不被當短 segment。
    回傳第一個穩定/合理段的起始 offset，否則 0。
    """
    p = Path(path)
    try:
        raw = np.fromfile(p, dtype="<i2", count=5000).astype(float)
        # 策略1: 原低 std 條件（適合有變異的 EEG/position）
        for i in range(0, len(raw) - 200, 5):
            seg = raw[i : i + 200]
            stdv = float(np.std(seg))
            mx = float(np.max(np.abs(seg)))
            if 5 < stdv < 400 and mx < 2000:
                return i
        # 策略2: 低變異但合理範圍資料起點（oximeter spo2/pleth/pulse 關鍵；值多在 ±2000 內）
        for i in range(0, len(raw) - 64, 4):
            seg = raw[i : i + 64]
            if np.all(np.abs(seg) < 2000):
                # 再看後 64 筆也合理，確認已過渡到真實資料
                if i + 128 < len(raw) and np.all(np.abs(raw[i + 64 : i + 128]) < 2000):
                    return i
        return 0
    except Exception:
        return 0


def _parse_ndf_channel_meta(path: PathLike) -> Dict:
    """從 .ndf 檔頭的 embedded channel header (UTF-16 混 binary) 解析真實 fs / label / unit。
    這是 spo2 / pleth / pulse 等 oximeter 通道最可靠的來源（SETUP.INI 常缺或靠易錯的 partial match）。
    同時嘗試提供 data_start 提示以利之後校正 samples 與 header skip。
    回傳例: {'fs': 3.0, 'label': 'SpO2', 'unit': '%', 'scale': 0.1, 'data_start_hint': 1316}
    找不到或失敗回傳 {} 。
    """
    p = Path(path)
    res: Dict = {}
    try:
        b = p.read_bytes()[:8192]
        t = b.decode("utf-16-le", errors="ignore")
        m = re.search(r"<Channel>(.*?)</Channel>", t, re.S | re.I)
        if m:
            xml = m.group(1)
            for tag, key in [
                ("SamplingRate", "fs"),
                ("Unit", "unit"),
                ("Label", "label"),
                ("Scale", "scale"),
                ("Type", "type"),
                ("Format", "format"),
            ]:
                mm = re.search(r"<" + tag + r">([^<]+)</" + tag + r">", xml, re.I)
                if mm:
                    val = mm.group(1).strip()
                    if key == "fs":
                        try:
                            res[key] = float(val)
                        except Exception:
                            pass
                    elif key in ("scale",):
                        try:
                            res[key] = float(val)
                        except Exception:
                            pass
                    else:
                        res[key] = val
        # 由 <Format> 決定樣本 dtype / bytes（int16 以外的通道若沿用 int16 會數值錯誤 + duration 倍增）
        fmt = str(res.get("format", "")).strip().lower()
        if fmt:
            dt, bps = _NDF_FORMAT_DTYPE.get(fmt, _DEFAULT_NDF_DTYPE)
            res["np_dtype"] = dt
            res["bytes_per_sample"] = bps

        # 結構化定位資料區起點：anchor(02 02 08 00 00 00) -> float64 取樣率 -> 01 02 + len -> 樣本資料。
        # 這比舊的 _detect_header_offset 啟發式精確（不會誤砍開頭數秒），且提供權威 float64 fs。
        ai = b.find(_NDF_DATA_ANCHOR)
        if ai >= 0:
            data_start = ai + _NDF_ANCHOR_TO_DATA
            if data_start < len(b) or data_start <= p.stat().st_size:
                res["data_start_byte"] = data_start
            # anchor 後緊接 float64 取樣率（較 <SamplingRate> 標籤更精確；通常一致）
            try:
                fs_f64 = float(np.frombuffer(b, dtype="<f8", count=1, offset=ai + 6)[0])
                if fs_f64 > 0:
                    res["fs_f64"] = fs_f64
                    # 僅在與標籤一致（<1% 差異）時採用 f64，避免異常 anchor 造成離譜 fs
                    tag_fs = res.get("fs")
                    if tag_fs is None or abs(fs_f64 - tag_fs) <= max(0.05, tag_fs * 0.01):
                        res["fs"] = fs_f64
            except Exception:
                pass

        # 解析嵌入的絕對起始時間（</Channel> 後的 UTF-16 ISO timestamp，如 20260511T220305.264000）。
        # 這是該通道「第一筆 sample」的絕對牆鐘時間；各通道因感測器分批上線而不同（D20 實測 spread 達 118s）。
        # 同時掃 latin1 視角以防偶發單位元組對齊差異。
        ts_m = re.search(r"(20\d{6})T(\d{6})(\.\d+)?", t) or re.search(
            r"(20\d{6})T(\d{6})(\.\d+)?", b.decode("latin1", errors="ignore")
        )
        if ts_m:
            try:
                dt0 = datetime.strptime(ts_m.group(1) + ts_m.group(2), "%Y%m%d%H%M%S")
                if ts_m.group(3):
                    dt0 = dt0 + timedelta(seconds=float(ts_m.group(3)))
                res["start_datetime"] = dt0
            except Exception:
                pass

        # 備援（合成檔 / 非 Nox 結構，無 anchor）：沿用 </Channel> 後的保守 hint，交給 _detect_header_offset
        if "data_start_byte" not in res:
            ch_end_pat = b"<\x00/\x00C\x00h\x00a\x00n\x00n\x00e\x00l\x00>\x00"
            ch_end = b.find(ch_end_pat)
            if ch_end > 0:
                res["data_start_hint"] = ch_end + 200
            else:
                ts_pat = b"2\x000\x002\x006"
                ts = b.find(ts_pat)
                if ts > 100:
                    res["data_start_hint"] = ts + 300
    except Exception:
        pass
    return res


class NoxRecording:
    """
    單一病患錄音的最佳化讀取器。
    開啟成本極低（僅解析 meta + memmap 標頭），讀取指定時間區段時才真正載入資料。
    """

    def __init__(self, patient_dir: PathLike):
        self.patient_dir = Path(patient_dir).resolve()
        self.name = self.patient_dir.name
        self.raw_dir = find_raw_data_dir(self.patient_dir)
        if self.raw_dir is None:
            raise FileNotFoundError(
                f"找不到 raw data 目錄（含 SETUP.INI + .ndf）：{self.patient_dir}"
            )

        self.setup_meta = parse_setup_ini(self.raw_dir / "SETUP.INI")
        self.device_info = self._parse_device_ini()
        self.ch_index = build_channel_index(self.raw_dir, self.setup_meta)
        self.duration_sec = compute_overall_duration(self.ch_index)
        self.start_datetime = self._guess_start_time()

        print(f"[NoxRecording] Init {self.name}: NDF channels={len(self.ch_index)}, raw_dur~{self.duration_sec:.1f}s")

        # 嘗試合併 EDF 中存在但 NDF 沒有的有效通道（按需求：先 NDF，再補 EDF only）
        self.edf_path = self._find_companion_edf()
        self._nox_edf_header: Optional[NoxEdfHeader] = None
        self._edf_uses_pyedflib = False
        if self.edf_path:
            print(f"[EDF] Companion found, will attempt merge (NDF priority for overlaps)")
            self._merge_edf_only_channels()
            print(f"[NoxRecording] After EDF merge: total ch_index keys={len(self.ch_index)} (some are aliases)")
        else:
            print("[EDF] No companion EDF for this recording.")

        # PR1: 套用 header_offset 偵測 + 低頻 fs 推斷（position 等 fs=0 問題 critical 修正）
        self._apply_header_offsets()
        self._infer_lowfreq_fs()
        # 刷新 duration_sec（Issue 2 修正）：infer 後部分 entry duration 更新，需重算 overall 以處理全低頻或 edge 情況
        self.duration_sec = compute_overall_duration(self.ch_index) or self.duration_sec
        self._apply_edf_duration_fallback()

        # B1: 計算全域時間原點與逐通道 offset（各通道因感測器分批上線而起始時間不同，須對齊到同一絕對軸）
        self.total_span_sec = self.duration_sec
        self._compute_time_origin_and_offsets()

        # 快取已開啟的 memmap（key = 正規化名稱）
        self._mmap_cache: Dict[str, np.memmap] = {}

        # EDF 來源絕對時間校準：
        # 優先已用 Data.ndb 的精確 RecordingStart 當 origin（次秒精確、EDF offset=0 即正確），無需再校準。
        # 僅當沒有 Data.ndb（origin 退回 EDF 整秒 header）時，才以 NDF 互相關估計 EDF 偏移作為備援。
        if not getattr(self, "_origin_from_ndb", False):
            try:
                self._calibrate_edf_offset()
            except Exception as ex:
                print(f"[EDF-cal] 互相關校準失敗（不影響載入）: {ex}")

    def _calibrate_edf_offset(self) -> None:
        """以 NDF 校準 EDF 來源通道的全域 offset（修正 EDF header 次秒遺失造成的整體偏移）。
        作法：找一個強訊號的 EDF/NDF 重疊通道，互相關量出 EDF↔NDF 偏移，
        套用到所有 EDF 來源 entry 的 start_offset_sec。無重疊（純 EDF）則不動。
        """
        # 收集同時有 EDF 與 NDF 變體的 base（dup pairs）
        edf_by_base: Dict[str, Dict] = {}
        ndf_by_base: Dict[str, Dict] = {}
        for k, e in self.ch_index.items():
            base = e.get("base_name")
            if not base:
                continue
            bl = base.lower()
            if e.get("source") == "edf":
                edf_by_base.setdefault(bl, e)
            else:
                ndf_by_base.setdefault(bl, e)
        common = [b for b in edf_by_base if b in ndf_by_base]
        if not common:
            return  # 無重疊通道，無法校準（純 EDF 錄音維持原樣）

        # 選校準通道：偏好連續訊號（呼吸/EEG/氣流），其次高 fs（互相關較穩）
        pref = ("flow", "c3", "c4", "nasal", "thermistor", "rip", "abdomen",
                "thorax", "ecg", "pleth")

        def score(b):
            for i, kw in enumerate(pref):
                if kw in b:
                    return (1000 - i, float(edf_by_base[b].get("fs", 0) or 0))
            return (0, float(edf_by_base[b].get("fs", 0) or 0))

        common.sort(key=score, reverse=True)

        lag = None
        chosen = None
        for b in common[:5]:
            lag = self._measure_edf_ndf_lag(edf_by_base[b], ndf_by_base[b])
            if lag is not None:
                chosen = b
                break
        if lag is None:
            print("[EDF-cal] 無法量測 EDF↔NDF 偏移（訊號不足/相關過低），維持 offset=0")
            return

        edf_off = -lag  # lag = EDF 相對 NDF；EDF 偏早(lag<0) → 需 +offset 往後移
        applied = 0
        for k, e in self.ch_index.items():
            if e.get("source") == "edf":
                e["start_offset_sec"] = edf_off
                applied += 1
        # 重算概觀右界（offset 變動後）
        span = 0.0
        seen_id = set()
        for e in self.ch_index.values():
            if id(e) in seen_id:
                continue
            seen_id.add(id(e))
            dur = float(e.get("duration_sec", 0.0) or 0.0)
            off = float(e.get("start_offset_sec", 0.0) or 0.0)
            if dur > 0:
                span = max(span, off + dur)
        if span > 0:
            self.total_span_sec = span
        print(f"[EDF-cal] 以「{chosen}」校準：EDF 來源 offset = {edf_off:+.3f}s "
              f"（lag={lag:+.3f}s, 套用 {applied} keys）")

    def _measure_edf_ndf_lag(self, edf_e: Dict, ndf_e: Dict) -> Optional[float]:
        """互相關量 EDF vs NDF 在全域時間的偏移（秒）。lag<0 表示 EDF 被擺早。
        於兩者都覆蓋的中段取 ~200s 窗，包絡降到 50Hz 共同格做互相關；相關過低回 None。"""
        fs_e = float(edf_e.get("inferred_fs") or edf_e.get("fs") or 0)
        fs_n = float(ndf_e.get("inferred_fs") or ndf_e.get("fs") or 0)
        if fs_e <= 0 or fs_n <= 0:
            return None
        off_e = float(edf_e.get("start_offset_sec", 0.0) or 0.0)
        off_n = float(ndf_e.get("start_offset_sec", 0.0) or 0.0)
        dur_e = float(edf_e.get("duration_sec", 0) or 0)
        dur_n = float(ndf_e.get("duration_sec", 0) or 0)
        g_lo = max(off_e, off_n) + 100.0
        g_hi = min(off_e + dur_e, off_n + dur_n) - 100.0
        if g_hi - g_lo < 80.0:
            return None
        g0 = g_lo + (g_hi - g_lo) * 0.4
        span = min(200.0, g_hi - g0)
        if span < 60.0:
            return None
        try:
            de, _ = self.get_data(edf_e["display_name"], start_sec=g0 - off_e,
                                  duration=span, as_float=True)
            dn, _ = self.get_data(ndf_e["display_name"], start_sec=g0 - off_n,
                                  duration=span, as_float=True)
        except Exception:
            return None
        if len(de) < 100 or len(dn) < 100:
            return None
        grid_fs = 50.0
        grid_t = np.arange(0.0, span, 1.0 / grid_fs)

        def envel(d, fs):
            d = np.abs(np.asarray(d, dtype=np.float64) - np.median(d))
            win = max(1, int(round(fs / grid_fs)))
            sm = np.convolve(d, np.ones(win) / win, mode="same")
            t = np.arange(len(d)) / fs
            return np.interp(grid_t, t, sm)

        ee = envel(de, fs_e)
        en = envel(dn, fs_n)
        m = min(len(ee), len(en))
        ee = ee[:m] - ee[:m].mean()
        en = en[:m] - en[:m].mean()
        if np.sum(ee ** 2) <= 0 or np.sum(en ** 2) <= 0:
            return None
        corr = np.correlate(ee, en, mode="full")
        lags = (np.arange(corr.size) - (m - 1)) / grid_fs
        mask = np.abs(lags) <= 5.0  # 合理偏移上限
        if not np.any(mask):
            return None
        sub = corr[mask]
        ci = int(np.argmax(sub))
        best_lag = float(lags[mask][ci])
        denom = (np.sqrt(np.sum(ee ** 2)) * np.sqrt(np.sum(en ** 2))) or 1.0
        if float(sub[ci] / denom) < 0.6:
            return None
        return best_lag

    def _edf_start_datetime(self) -> Optional[datetime]:
        """取得 companion EDF 的起始時間（Nox 專有 parser 或標準 pyedflib），作為全域原點首選。"""
        if not self.edf_path:
            return None
        if self._nox_edf_header is not None and self._nox_edf_header.start_datetime:
            return self._nox_edf_header.start_datetime
        _, edf_start = _read_standard_edf_meta(self.edf_path)
        if edf_start is not None:
            return edf_start
        try:
            return parse_nox_edf(self.edf_path).start_datetime
        except Exception:
            return None

    def _read_ndb_recording_start(self) -> Optional[datetime]:
        """從 raw_dir/Data.ndb 讀取權威錄音起始時間（含次秒）。
        Nox Noxturnal 把精確起始存在 internal_property 的 RecordingInfo/RecordingStart，
        型別 'Ticks' = .NET DateTime ticks（100ns since 0001-01-01）。
        EDF header 的 starttime 只存整秒、會捨去次秒（實測 D18 = 0.744s）；此值才是精確原點。
        無 Data.ndb 或讀取失敗回 None。"""
        try:
            raw = getattr(self, "raw_dir", None)
            if not raw:
                return None
            ndb = Path(raw) / "Data.ndb"
            if not ndb.exists():
                return None
            conn = sqlite3.connect(f"file:{ndb}?mode=ro&immutable=1", uri=True)
            try:
                row = conn.execute(
                    "SELECT value FROM internal_property "
                    "WHERE name='RecordingInfo' AND key='RecordingStart' AND type='Ticks' LIMIT 1"
                ).fetchone()
            finally:
                conn.close()
            if not row or row[0] in (None, ""):
                return None
            ticks = int(row[0])
            dt = datetime(1, 1, 1) + timedelta(microseconds=ticks / 10)
            print(f"[NDB] RecordingStart（權威精確起始）= {dt}（次秒 {ticks % 10_000_000 / 1e7:.3f}s）")
            return dt
        except Exception as e:
            print(f"[NDB] 讀取 RecordingStart 失敗（改用 EDF/NDF 推估）: {e}")
            return None

    def _compute_time_origin_and_offsets(self) -> None:
        """建立統一絕對時間軸：
        - 原點 origin = EDF 起始時間（最權威，且實測 ≈ 最早 NDF 通道、略早於它）；無 EDF 則取最早 NDF start。
        - 各通道 start_offset_sec = max(0, start_datetime - origin)；EDF 來源 = 0（與 origin 同鐘）。
        - self.start_datetime = origin；self.total_span_sec = max(offset + duration)（概觀軸右界）。
        舊行為（所有通道 offset=0）只在完全沒有任何嵌入時間時保留。
        """
        # 收集唯一通道（以 path 去重，避免 alias 重複計入）的 NDF 起始時間
        seen_paths = set()
        ndf_starts: List[datetime] = []
        for e in self.ch_index.values():
            if e.get("source") == "edf":
                continue
            pstr = str(e.get("path", ""))
            if pstr in seen_paths:
                continue
            seen_paths.add(pstr)
            sd = e.get("start_datetime")
            if isinstance(sd, datetime):
                ndf_starts.append(sd)

        edf_start = self._edf_start_datetime()
        ndb_start = self._read_ndb_recording_start()  # 權威精確起始（含次秒）

        # 決定原點：
        # 1) Data.ndb 的 RecordingStart（.NET ticks，含次秒，最權威）。這是整個錄音 session 的 t=0，
        #    NDF/EDF/評分標記同屬此 session 共用之。EDF header 只存整秒會遺失次秒（D18 實測 floor 比
        #    真實起始早 0.744s，正是 EDF 波形偏移來源）。
        #    **安全核對**：有 companion EDF 時，floor(RecordingStart) 必須等於 EDF header 起始（證明
        #    Data.ndb 與此 EDF 為同一次錄音/匯出），不符則不採信精確值、退回 EDF header。
        # 2) 否則 EDF header 起始（整秒）；3) 否則最早 NDF；都沒有則維持舊 start_datetime 全 offset=0。
        self._origin_from_ndb = False
        ndb_ok = ndb_start is not None and (
            edf_start is None
            or abs((ndb_start.replace(microsecond=0) - edf_start).total_seconds()) <= 1.0
        )
        if ndb_start is not None and not ndb_ok:
            print(f"[NDB] RecordingStart({ndb_start}) 與 EDF header({edf_start}) 不一致，"
                  f"不採用精確值（疑非同次匯出）")
        origin: Optional[datetime] = None
        if ndb_ok:
            origin = ndb_start
            self._origin_from_ndb = True
        elif edf_start is not None:
            origin = edf_start
            if ndf_starts:
                # 防呆：若某 NDF 早於 EDF start（實測未見，但保險），原點取更早者以免出現負 offset
                origin = min([edf_start] + ndf_starts)
        elif ndf_starts:
            origin = min(ndf_starts)

        if origin is None:
            # 無任何嵌入時間（如合成測試檔）：維持舊語意，全部對齊 0
            self._time_origin = self.start_datetime
            span = 0.0
            for e in self.ch_index.values():
                e["start_offset_sec"] = 0.0
                span = max(span, float(e.get("duration_sec", 0.0) or 0.0))
            self.total_span_sec = span or self.duration_sec
            return

        self._time_origin = origin
        self.start_datetime = origin

        span = 0.0
        for e in self.ch_index.values():
            dur = float(e.get("duration_sec", 0.0) or 0.0)
            if e.get("source") == "edf":
                off = 0.0
            else:
                sd = e.get("start_datetime")
                off = max(0.0, (sd - origin).total_seconds()) if isinstance(sd, datetime) else 0.0
            e["start_offset_sec"] = off
            if dur > 0:
                span = max(span, off + dur)

        # EDF 整體時長通常涵蓋最久（end ≥ NDF max-end），併入概觀右界
        if self.edf_path:
            edf_dur = _edf_duration_sec(self.edf_path)
            if edf_dur > span:
                span = edf_dur
        self.total_span_sec = span or self.duration_sec

    def channel_offset_sec(self, name: str) -> float:
        """回傳通道相對全域原點的時間偏移（秒）；供 viewer 在絕對時間軸上正確擺放波形。"""
        info = self.get_channel_info(name)
        if not info:
            return 0.0
        return float(info.get("start_offset_sec", 0.0) or 0.0)

    def _apply_edf_duration_fallback(self) -> None:
        """NDF 推算時長不可靠時（例如僅有 rip phase.ndf），改採 companion EDF 時長。"""
        if self.duration_sec >= _MIN_RELIABLE_NDF_DURATION_SEC:
            return
        if not self.edf_path:
            return
        edf_dur = _edf_duration_sec(self.edf_path)
        if edf_dur < _MIN_RELIABLE_NDF_DURATION_SEC:
            return
        print(
            f"[NoxRecording] NDF duration {self.duration_sec:.1f}s unreliable, "
            f"using EDF duration {edf_dur:.1f}s"
        )
        self.duration_sec = edf_dur
        if self.start_datetime is None:
            _, edf_start = _read_standard_edf_meta(self.edf_path)
            if edf_start is not None:
                self.start_datetime = edf_start

    def _parse_device_ini(self) -> Dict:
        """
        手動 section 解析 DEVICE.INI，只取具 key=val 的可靠區塊（如 [DeviceInfo]）。
        避開 [Channels] 純 list 格式（"Light 1Hz lx" 等無 =），否則 configparser 失敗導致 device_info={}。
        """
        ini_path = self.raw_dir / "DEVICE.INI"
        info: Dict = {}
        if not ini_path.exists():
            return info
        try:
            content = ini_path.read_text(encoding="utf-8", errors="ignore")
            lines = content.splitlines()
            i = 0
            while i < len(lines):
                line = lines[i].strip()
                if line.startswith("[") and line.endswith("]"):
                    sect = line[1:-1]
                    sect_dict = {}
                    j = i + 1
                    while j < len(lines):
                        l = lines[j].strip()
                        if l.startswith("[") and l.endswith("]"):
                            break
                        if "=" in l:
                            k, v = l.split("=", 1)
                            sect_dict[k.strip()] = v.strip()
                        j += 1
                    if sect_dict:  # 只保留有 key=val 的 section，跳過 [Channels] list
                        info[sect] = sect_dict
                    i = j
                    continue
                i += 1
        except Exception:
            pass
        return info

    def _guess_start_time(self) -> Optional[datetime]:
        # 從資料夾名稱或 raw 資料夾名稱猜（例如 20260511T220227）
        folder = self.raw_dir.name
        m = re.search(r"(\d{8})T(\d{6})", folder)
        if m:
            try:
                return datetime.strptime(m.group(1) + m.group(2), "%Y%m%d%H%M%S")
            except Exception:
                pass
        # 嘗試從 edf 檔名（使用與 _find 相同的多位置搜尋，避免 patient_dir 為 raw 層時漏掉同級 EDF）
        for edf in self._find_edf_candidates():
            m = re.search(r"(\d{8})", edf.name)
            if m:
                try:
                    return datetime.strptime(m.group(1), "%Y%m%d")
                except Exception:
                    pass
        return None

    def _find_edf_candidates(self) -> list[Path]:
        """收集可能的 companion EDF 候選。
        搜尋位置：
        - patient_dir 本身（最常見：EDF 與 Dxx_xxx 資料夾同級）
        - raw_dir 的 parent（明確支援「EDF 跟 raw data 資料夾同級」的情境）
        - raw_dir 本身（以防萬一 EDF 被放在 raw 裡）
        會去重，並印出實際搜尋的目錄，讓使用者從 console 明確看到「現在從哪裡找」。
        """
        search_dirs: list[Path] = []
        if self.patient_dir and self.patient_dir.exists():
            search_dirs.append(self.patient_dir)
        if getattr(self, "raw_dir", None):
            rp = self.raw_dir.parent
            if rp and rp.exists() and rp != self.patient_dir:
                search_dirs.append(rp)
            if self.raw_dir.exists():
                search_dirs.append(self.raw_dir)
        # 也往上找一層（萬一 patient_dir 指向太深）
        if self.patient_dir:
            up = self.patient_dir.parent
            if up and up.exists() and up != self.patient_dir and up not in search_dirs:
                search_dirs.append(up)

        print(f"[EDF] _find_companion_edf: patient_dir={self.patient_dir}")
        if getattr(self, "raw_dir", None):
            print(f"[EDF] raw_dir={self.raw_dir}")
        print(f"[EDF] 將在以下位置搜尋 *.edf（同級 raw data 也會被涵蓋）： {[str(s) for s in search_dirs]}")

        seen = set()
        cands: list[Path] = []
        for sd in search_dirs:
            try:
                for p in sd.glob("*.edf"):
                    if p.is_file():
                        rp = p.resolve()
                        if rp not in seen:
                            seen.add(rp)
                            cands.append(p)
            except Exception:
                pass
        print(f"[EDF] 在以上位置共找到 {len(cands)} 個 .edf 候選: {[str(c) for c in cands]}")
        return cands

    def _find_companion_edf(self) -> Optional[Path]:
        """在 patient_dir 及其同級（raw data 旁）尋找伴隨的 .edf 檔案（裝置匯出的 EDF）。
        優先選擇檔名含 'edf' 且檔案較大者（通常是主 EDF）。
        只回傳路徑，不在此檢查是否可開啟（開啟失敗會在 merge 時被 catch）。
        現在會同時搜 patient_dir + raw data 資料夾的同級目錄，確保「EDF 跟 raw data 資料夾同級」時能找到。
        """
        if not self.patient_dir or not self.patient_dir.exists():
            print("[EDF] No patient_dir")
            return None
        cands = self._find_edf_candidates()
        if not cands:
            return None
        # 排序：優先檔名含 edf，其次大小最大
        cands.sort(key=lambda p: (0 if "edf" in p.name.lower() else 1, -p.stat().st_size))
        chosen = cands[0]
        print(f"[EDF] Chosen companion EDF: {chosen}")
        return chosen

    def _prepare_edf_reader(self, edf_path: Path) -> None:
        """準備 EDF 讀取：標準 EDF 走 pyedflib，Nox 專有格式走自訂 parser。"""
        if self._nox_edf_header is not None or self._edf_uses_pyedflib:
            return
        ok, msg = try_pyedflib_open(edf_path)
        if ok:
            print(f"[EDF] Standard EDF open OK via pyedflib: {msg}")
            self._edf_uses_pyedflib = True
            return
        print(f"[EDF] pyedflib cannot open ({msg}); using Nox proprietary EDF parser")
        try:
            self._nox_edf_header = parse_nox_edf(edf_path)
            print(
                f"[EDF] Nox parser OK: n_signals={self._nox_edf_header.n_signals}, "
                f"n_records={self._nox_edf_header.n_records}, "
                f"duration={self._nox_edf_header.duration_sec:.1f}s"
            )
        except Exception as e:
            print(f"[EDF] Nox parser failed for {edf_path}: {e}")
            raise

    def _merge_edf_only_channels(self):
        """解析 EDF 中存在、但 NDF ch_index 中沒有的有效通道，並合併進 ch_index。
        Nox 專有 EDF 使用自訂 parser；標準 EDF 仍走 pyedflib。
        """
        if not self.edf_path or not Path(self.edf_path).exists():
            print("[EDF] _merge called but no valid edf_path, skip merge")
            return
        print(f"[EDF] Starting _merge_edf_only_channels for {self.edf_path}")
        try:
            self._prepare_edf_reader(self.edf_path)
        except Exception as e:
            print(f"[EDF] FATAL: cannot prepare EDF reader: {e}")
            self.edf_path = None
            return

        # 既有 NDF 通道：建立 正規化名 → 唯一 NDF entry 對照（用於重疊偵測 + 重貼標）。
        existing_norms = set()
        ndf_by_norm: Dict[str, Dict] = {}
        for k, e in list(self.ch_index.items()):
            stem = e.get("stem", k)
            existing_norms.add(_normalize(stem))
            existing_norms.add(_normalize(k))
            if e.get("source") != "edf":
                ndf_by_norm.setdefault(_normalize(stem), e)

        added = 0          # EDF-only 新增數
        dup = 0            # 與 NDF 重疊、改為並列(EDF/NDF 後綴)的數量
        added_labels = []

        if self._edf_uses_pyedflib:
            try:
                import pyedflib

                f = pyedflib.EdfReader(str(self.edf_path))
                n_sig = f.signals_in_file
                nsamp_all = f.getNSamples()
                for i in range(n_sig):
                    label = (f.getLabel(i) or "").strip()
                    if not label:
                        continue
                    fs = float(f.getSampleFrequency(i) or 0)
                    if fs <= 0:
                        continue
                    unit = (f.getPhysicalDimension(i) or "").strip()
                    ns = int(nsamp_all[i]) if i < len(nsamp_all) else 0
                    entry = self._make_edf_channel_entry(label, fs, unit, ns, i)
                    kind = self._merge_one_edf_signal(entry, existing_norms, ndf_by_norm)
                    if kind == "dup":
                        dup += 1
                    elif kind == "added":
                        added += 1
                    if kind:
                        added_labels.append(f"{label}(fs={fs:.2f})")
                f.close()
            except Exception as e:
                print(f"[EDF] pyedflib merge failed: {e}")
                self.edf_path = None
                return
        elif self._nox_edf_header:
            hdr = self._nox_edf_header
            for i, label in enumerate(hdr.labels):
                if not label:
                    continue
                fs = hdr.fs_for_channel(i)
                if fs <= 0:
                    continue
                _, unit = _guess_fs_unit_from_label(label)
                ns = hdr.total_samples_for_channel(i)
                entry = self._make_edf_channel_entry(label, fs, unit, ns, i)
                entry["nox_edf"] = True
                kind = self._merge_one_edf_signal(entry, existing_norms, ndf_by_norm)
                if kind == "dup":
                    dup += 1
                elif kind == "added":
                    added += 1
                if kind:
                    added_labels.append(f"{label}(fs={fs:.2f})")

        if added or dup:
            self.duration_sec = compute_overall_duration(self.ch_index) or self.duration_sec
            print(f"[EDF] Merged: EDF-only={added}, 與NDF重疊並列(EDF/NDF)={dup}. "
                  f"duration={self.duration_sec:.1f}s")
        else:
            print("[EDF] No EDF channels merged (all invalid)")

    def _merge_one_edf_signal(self, entry: Dict, existing_norms: set,
                              ndf_by_norm: Dict[str, Dict]) -> Optional[str]:
        """處理單一 EDF signal entry：
        - 與 NDF 同名(正規化)重疊 → 兩者都保留並加 (EDF)/(NDF) 後綴顯示名；回傳 'dup'。
        - 否則 EDF-only → 照常註冊（plain 名）；回傳 'added'。
        """
        label = entry["stem"]
        lnorm = _normalize(label)
        nd = ndf_by_norm.get(lnorm)
        if nd is not None:
            # 重疊：兩個變體共用 NDF stem 當 base（與事件名/歷史比對一致，避免 EDF 標籤大小寫差異
            # 讓嚴格事件比對失配）。EDF 變體仍 source=edf、路由 EDF 資料。
            nd_stem = nd.get("stem", label)
            disp_edf = f"{nd_stem}(EDF)"
            entry["display_name"] = disp_edf
            entry["base_name"] = nd_stem
            self.ch_index[disp_edf] = entry
            self.ch_index[_normalize(disp_edf)] = entry
            # 對應 NDF 變體重貼標為 (NDF)（只做一次）
            if not nd.get("display_name"):
                disp_ndf = f"{nd_stem}(NDF)"
                nd["display_name"] = disp_ndf
                nd["base_name"] = nd_stem
                self.ch_index[disp_ndf] = nd
                self.ch_index[_normalize(disp_ndf)] = nd
            return "dup"
        # EDF-only
        if lnorm in existing_norms:
            return None  # 與既有 EDF entry 重複（罕見），略過
        self._register_edf_channel(entry, existing_norms)
        return "added"

    def _make_edf_channel_entry(
        self, label: str, fs: float, unit: str, samples: int, sig_index: int
    ) -> Dict:
        dur = samples / fs if fs > 0 else 0.0
        return {
            "path": str(self.edf_path),
            "stem": label,
            "fs": fs,
            "unit": unit,
            "desc": label,
            "internal_key": None,
            "samples": samples,
            "duration_sec": dur,
            "header_offset": 0,
            "inferred_fs": None,
            "source": "edf",
            "edf_signal_index": sig_index,
        }

    def _register_edf_channel(self, entry: Dict, existing_norms: set) -> None:
        label = entry["stem"]
        lnorm = _normalize(label)
        if label not in self.ch_index:
            self.ch_index[label] = entry
        if lnorm and lnorm not in self.ch_index:
            self.ch_index[lnorm] = entry
        existing_norms.add(lnorm)

    def _apply_header_offsets(self):
        """對 ch_index 內每個唯一 .ndf 套用 _detect_header_offset（所有 .ndf 均套用；D18/D20 實測 junk prefix 普遍存在於 position + spo2 + c3 等）。
        同時校正 samples（扣除 header 解讀 phantom），確保 spo2 等低頻通道的 duration 與 sample count 正確（避免被當成只有幾分鐘的短 segment）。
        EDF 來源通道跳過（header_offset=0）。
        """
        seen = set()
        for e in self.ch_index.values():
            if e.get("source") == "edf":
                e["header_offset"] = 0
                continue
            # 結構化通道：data_start_byte 已精確定位資料起點，samples/duration 已正確，跳過啟發式偵測
            if e.get("data_start_byte") is not None:
                e["header_offset"] = 0
                continue
            pstr = str(e.get("path", ""))
            if pstr in seen:
                continue
            seen.add(pstr)
            try:
                off = _detect_header_offset(e["path"])
                hint = e.get("_data_start_hint")
                if hint:
                    hint_off = max(0, hint // 2)
                    if hint_off > off:
                        off = hint_off
                e["header_offset"] = off
                view_samp = e.get("samples", 0) or 0
                if off > 0 and view_samp > off and e.get("fs", 0) > 0:
                    e["duration_sec"] = (view_samp - off) / e["fs"]
            except Exception:
                e["header_offset"] = 0

    def _infer_lowfreq_fs(self):
        """
        低頻 fs inference（critical）：position-like (stem=="position" 或 fs==0 + 大 samples)
        從 overall duration 反推 fs，或對 position 使用 ~20Hz。
        存 "inferred_fs" 並更新 "fs"；get_data 內 offset 總以 sample 單位。
        EDF 來源跳過（已有正確 fs）。
        """
        overall = self.duration_sec or 0.0
        for e in self.ch_index.values():
            if e.get("source") == "edf":
                continue
            # 結構化通道：fs 來自 header float64（如 position 20.0002），已正確，不可用 samples/overall 反推覆蓋
            if e.get("data_start_byte") is not None:
                continue
            stem = str(e.get("stem", "")).lower()
            fs = e.get("fs", 0.0) or 0.0
            samples = e.get("samples", 0) or 0
            is_pos_like = stem == "position" or (
                fs <= 0
                and samples > 10000
                and "axis" not in stem
                and "impedance" not in stem
                and "imp" not in stem
            )
            if is_pos_like:
                if overall > 60 and samples > 0:
                    inf = samples / overall
                    e["inferred_fs"] = round(inf, 4)
                    e["fs"] = e["inferred_fs"]
                    if e["fs"] > 0:
                        e["duration_sec"] = samples / e["fs"]
                elif stem == "position":
                    e["inferred_fs"] = 20.0
                    e["fs"] = 20.0
                    e["duration_sec"] = samples / 20.0 if samples > 0 else 0.0

    def is_position_channel(self, name: str) -> bool:
        """判斷是否 position 類別通道（僅 stem 或 name/desc 含 "position"；移除寬鬆低 fs 判斷以避免誤判 pulse / *impedance / set pressure 等）。"""
        info = self.get_channel_info(name)
        if not info:
            return False
        nm = _normalize(name)
        stem = str(info.get("stem", "")).lower()
        desc = str(info.get("desc", "")).lower()
        return (
            stem == "position"
            or "position" in nm
            or "position" in desc
        )

    def get_patient_info(self, mask_pii: bool = False) -> Dict:
        """
        使用 python-docx 從 word_*.docx 解析 demographics + clinical tables。
        特別提取 Age/Gender/BMI、BODY POSITION SUMMARY / RDI per posture。
        搜尋 patient_dir 優先 *word*.docx，fallback raw_dir。
        無 docx 時 fallback 顯示技術資訊。
        回傳: {"demographics": {...}, "clinical": {"rdi":.., "position_summary": [...] , ...}, "source": ...}
        """
        info = {"demographics": {}, "clinical": {}, "source": None}
        candidates = []
        for base in (self.patient_dir, self.raw_dir):
            if base and base.exists():
                for pat in ("*word*.docx", "*PSG*.docx", "*.docx"):
                    for f in base.glob(pat):
                        if f.is_file():
                            candidates.append(f)
        if not candidates:
            info["demographics"] = {"mrn": self.name, "note": "報告未找到，僅顯示技術資訊"}
            return info
        # 選最佳（含 word 或 match name）
        docx_path = None
        for c in candidates:
            if "word" in c.name.lower() or self.name[:6] in c.name:
                docx_path = c
                break
        if docx_path is None:
            docx_path = candidates[0]
        info["source"] = str(docx_path)
        try:
            from docx import Document
            doc = Document(str(docx_path))
            demo = {}
            full_text = "\n".join(p.text for p in doc.paragraphs)
            # 從 paras 提取（Age:50\t , Gender:Male , Height:175.0 cm , Weight:90.7kg , BMI:29.6 , MR# , Recording Date）
            m = re.search(r"Age:\s*(\d+)", full_text, re.IGNORECASE)
            if m:
                demo["age"] = int(m.group(1))
            m = re.search(r"Gender:\s*(Male|Female|男|女)", full_text, re.IGNORECASE)
            if m:
                demo["gender"] = m.group(1)
            m = re.search(r"Height:\s*([\d.]+)", full_text, re.IGNORECASE)
            if m:
                demo["height_cm"] = float(m.group(1))
            m = re.search(r"Weight:\s*([\d.]+)", full_text, re.IGNORECASE)
            if m:
                demo["weight_kg"] = float(m.group(1))
            m = re.search(r"BMI:\s*([\d.]+)", full_text, re.IGNORECASE)
            if m:
                demo["bmi"] = float(m.group(1))
            m = re.search(r"MR#:\s*([A-Za-z0-9\-]+)", full_text, re.IGNORECASE)
            if m:
                demo["mrn"] = m.group(1)
            m = re.search(r"Recording Date:\s*([\d/]+)", full_text, re.IGNORECASE)
            if m:
                demo["study_date"] = m.group(1)
            m = re.search(r"Date of Birth:\s*([\d/]+)", full_text, re.IGNORECASE)
            if m:
                demo["dob"] = m.group(1)
            # 強化從 split paras
            for p in doc.paragraphs:
                t = p.text
                if "\t" in t or ":" in t:
                    if "Age:" in t and "age" not in demo:
                        try:
                            val = t.split("Age:")[1].split()[0].split("Y")[0].strip()
                            demo["age"] = int(val)
                        except Exception:
                            pass
                    if "BMI:" in t and "bmi" not in demo:
                        try:
                            val = t.split("BMI:")[1].split()[0].strip()
                            demo["bmi"] = float(val)
                        except Exception:
                            pass
            if not demo.get("mrn"):
                demo["mrn"] = self.name
            if mask_pii:
                if demo.get("mrn"):
                    demo["mrn"] = str(demo["mrn"])[:5] + "***"
                if demo.get("dob"):
                    demo["dob"] = "****"
            info["demographics"] = demo

            # clinical + position summary (Table with Supine/Left etc + RDI)
            clin = {}
            pos_sum = []
            seen_pos = set()
            for table in doc.tables:
                rtxt = [[c.text.strip().replace("\n", " ") for c in r.cells] for r in table.rows]
                # 僅處理有 'min' 或 'Time in Position' 的 BODY POSITION SUMMARY 類表格，避免 RDI-only 表重複
                is_body_pos = any(
                    "Time in Position" in " ".join(row)
                    or ("min" in " ".join(row).lower() and "RDI" in " ".join(row))
                    for row in rtxt[:2]
                )
                if is_body_pos or any(
                    row[0] in ("Supine", "Left", "Right", "Pron", "Up")
                    and "min" in " ".join(row).lower()
                    for row in rtxt
                ):
                    # 最小 robustness 改進：由 header 動態找 time/RDI col（而非硬 index 1/3），避免未來表格順序變動
                    header = rtxt[0] if rtxt else []
                    tcol = next((i for i, h in enumerate(header) if "time" in h.lower() or "min" in h.lower()), 1)
                    rdicol = next((i for i, h in enumerate(header) if "rdi" in h.lower()), 3)
                    for row in rtxt:
                        pos = row[0]
                        if pos in ("Supine", "Left", "Right", "Pron", "Up") and pos not in seen_pos:
                            try:
                                tmin = float(row[tcol]) if tcol < len(row) and row[tcol] else 0.0
                                rdi_str = row[rdicol] if rdicol < len(row) and row[rdicol] else ""
                                rdi = float(rdi_str) if rdi_str else 0.0
                                pos_sum.append(
                                    {"position": pos, "time_min": tmin, "rdi_per_hr": rdi}
                                )
                                seen_pos.add(pos)
                            except Exception:
                                pass
                # RDI from tables (只取一次)
                for row in rtxt:
                    j = " ".join(row)
                    if ("RDI" in j or "AHI" in j) and re.search(r"(\d+\.?\d*)\s*/?hr", j, re.I):
                        m = re.search(r"(\d+\.?\d*)\s*/?hr", j, re.I)
                        if m and "rdi" not in clin:
                            try:
                                clin["rdi"] = float(m.group(1))
                            except Exception:
                                pass
            if pos_sum:
                clin["position_summary"] = pos_sum
            # extra from paras (Total Recording etc)
            for p in doc.paragraphs:
                t = p.text
                if "Total Recording Time" in t:
                    m = re.search(r"([\d.]+)hrs", t)
                    if m:
                        clin["duration_h"] = float(m.group(1))
                if "Sleep Efficiency" in t:
                    m = re.search(r"([\d.]+)%", t)
                    if m:
                        clin["sleep_efficiency_pct"] = float(m.group(1))
            info["clinical"] = clin
        except ImportError:
            info["demographics"]["note"] = "報告解析失敗: python-docx 未安裝或 import 失敗"
            if not info["demographics"].get("mrn"):
                info["demographics"]["mrn"] = self.name
        except Exception as ex:
            info["demographics"]["note"] = f"報告解析失敗: {ex}"
            if not info["demographics"].get("mrn"):
                info["demographics"]["mrn"] = self.name
        if not info.get("demographics"):
            info["demographics"] = {"mrn": self.name, "note": "報告未找到，僅顯示技術資訊"}
        return info

    @property
    def channels(self) -> List[Dict]:
        """回傳所有獨特通道的摘要列表（以 stem 為主）。
        注意：'samples' 與 'duration_sec' 基於原始檔案 bytes//2（可能含 header junk bytes）；'header_offset' 已暴露可用來計算 effective samples = max(0, samples - header_offset)。對 fidelity 無影響（off << total）。
        """
        seen = set()
        out = []
        for k, e in self.ch_index.items():
            # 以「顯示名」去重（重疊通道的 NDF/EDF 變體各有唯一後綴名；非重疊則 = stem）。
            disp = e.get("display_name") or e["stem"]
            if disp in seen:
                continue
            seen.add(disp)
            out.append(
                {
                    "name": disp,
                    "base_name": e.get("base_name") or e["stem"],
                    "source": e.get("source", "ndf"),
                    "fs": e["fs"],
                    "unit": e["unit"],
                    "desc": e["desc"],
                    "samples": e["samples"],
                    "duration_sec": e["duration_sec"],
                    "path": str(e["path"]),
                    "header_offset": e.get("header_offset", 0),
                    "inferred_fs": e.get("inferred_fs"),
                    "source_label": format_channel_source(e),
                    "start_offset_sec": float(e.get("start_offset_sec", 0.0) or 0.0),
                    "start_datetime": e.get("start_datetime"),
                }
            )
        out.sort(key=lambda x: (-x["fs"], x["name"]))
        return out

    def list_channels(self, include_imp: bool = False) -> List[str]:
        """簡易列出人類可讀的通道名稱。"""
        names = []
        for c in self.channels:
            nm = c["name"]
            if not include_imp and "imp" in _normalize(nm):
                continue
            names.append(nm)
        return names

    def get_channel_info(self, name: str) -> Optional[Dict]:
        nk = _normalize(name)
        if nk in self.ch_index:
            return self.ch_index[nk].copy()
        # 模糊
        for k, e in self.ch_index.items():
            if _normalize(k) == nk or nk in _normalize(k):
                return e.copy()
        return None

    def _get_mmap(self, entry: Dict) -> np.memmap:
        key = str(entry["path"])
        if key in self._mmap_cache:
            return self._mmap_cache[key]
        dt = entry.get("np_dtype") or _DEFAULT_NDF_DTYPE[0]
        ds = entry.get("data_start_byte")
        if ds is not None:
            # 結構化路徑：自資料起點 byte、以 <Format> 對應 dtype、限定真實樣本數映射。
            # mm[0] 即第一筆真實樣本（offset 處理 header，且支援非 4-byte 對齊的 int32 起點）。
            n = int(entry.get("samples", 0) or 0)
            if n > 0:
                mm = np.memmap(entry["path"], dtype=dt, mode="r", offset=ds, shape=(n,))
            else:
                mm = np.memmap(entry["path"], dtype=dt, mode="r", offset=ds)
        else:
            # legacy：整檔映射（預設 int16），header junk 由 get_data 的 header_offset 切除
            mm = np.memmap(entry["path"], dtype=dt, mode="r")
        self._mmap_cache[key] = mm
        return mm

    def get_data(
        self,
        name: str,
        start_sec: float = 0.0,
        duration: Optional[float] = None,
        as_float: bool = True,
        *,
        strip_header: bool = True,
    ) -> Tuple[np.ndarray, float]:
        """
        讀取單一通道指定時間區段的資料。
        回傳 (data_array, fs)

        - 使用 memmap 切片，記憶體效率極高。
        - start_sec: 從錄音開始算起的秒數（以秒為單位，內部轉 sample）。
        - duration: 要讀取的秒數，None 表示到結尾。
        - strip_header: 預設 True，自動移除 .ndf 開頭 XML junk header（所有 .ndf；position 等典型 offset 5-80 samples，D18 pos=5, c3/spo2~50~650 視 header 大小；現在有改善偵測 + ndf header 解析）。
          偏移量以 sample 為單位套用：effective_start = round(start_sec * fs) + header_offset。
          即使 fs fallback 仍正確（使用 inferred_fs 優先）。
          設 False 則維持舊 raw 行為（含 header），供對照或相容。
        - 向後相容：無關鍵字參數呼叫預設 strip=True。
        - 所有 .ndf 均自動 strip（D18/D20 實測開頭 junk prefix 普遍存在，position/c3/spo2 等 offset 通常 5-80 samples，spo2 類因 header 結構較大；改善偵測 + ndf 解析後時間軸正確）。
        - 低頻道（position）使用 inferred_fs 確保 30s 視窗 sample count 正確（~ fs*30），時間軸精準。
        - export_to_standard_edf 使用 stripped 計算 pmin/pmax 更正確，避免 junk 污染 position 等特別受益。
        - 額外支援：若 patient_dir 旁有 .edf 檔案，會自動補充「NDF 中沒有、但 EDF 中有的有效通道」（讀取 pyedflib 物理值）。
          優先順序：先完整解析 NDF 通道，再合併 EDF-only 通道，一同顯示在 viewer 左側列表與波形中。
        """
        info = self.get_channel_info(name)
        if info is None:
            raise KeyError(f"找不到通道：{name}，可用：{self.list_channels()[:10]}...")

        # EDF-only 通道：直接從 pyedflib 讀物理值（readSignal 已 scaling）
        if info.get("source") == "edf":
            print(f"[EDF] get_data routed to EDF source for channel='{name}' (stem={info.get('stem')})")
            return self._get_data_edf(info, start_sec, duration, as_float)

        fs = info.get("inferred_fs") or info["fs"]
        if fs <= 0:
            fs = 1.0
        samples = info["samples"]
        # 結構化通道：memmap 已自資料起點映射（mm[0]=第一筆真實樣本），不再加 header_offset。
        # legacy 通道（無 anchor）：維持以 header_offset 切除開頭 junk 的舊行為。
        if info.get("data_start_byte") is not None:
            header_off = 0
        else:
            header_off = (info.get("header_offset", 0) or 0) if strip_header else 0

        start_sample = int(round(start_sec * fs)) + header_off
        if start_sample < 0:
            start_sample = 0
        if start_sample >= samples:
            return np.array([], dtype=np.float32 if as_float else np.int16), fs

        if duration is None or duration <= 0:
            end_sample = samples
        else:
            end_sample = start_sample + int(round(duration * fs))
            if end_sample > samples:
                end_sample = samples

        mm = self._get_mmap(info)
        raw = mm[start_sample:end_sample]
        if as_float:
            data = raw.astype(np.float32)
        else:
            data = raw
        return data, fs

    def _get_data_edf(
        self,
        info: Dict,
        start_sec: float = 0.0,
        duration: Optional[float] = None,
        as_float: bool = True,
    ) -> Tuple[np.ndarray, float]:
        """從伴隨 EDF 讀取指定通道（Nox 專有格式或標準 pyedflib）。"""
        edf_path = info.get("path")
        sig_idx = info.get("edf_signal_index")
        fs = info.get("fs") or 1.0
        if not edf_path or sig_idx is None:
            return np.array([], dtype=np.float32 if as_float else np.int16), fs

        total_samples = int(info.get("samples", 0))
        start_samp = max(0, int(round(start_sec * fs)))
        if duration is None or duration <= 0:
            n_samp = total_samples - start_samp
        else:
            n_samp = int(round(duration * fs))
        if start_samp >= total_samples or n_samp <= 0:
            return np.array([], dtype=np.float32 if as_float else np.int16), fs
        n_samp = min(n_samp, total_samples - start_samp)

        try:
            if info.get("nox_edf") or self._nox_edf_header:
                if self._nox_edf_header is None:
                    self._prepare_edf_reader(Path(edf_path))
                hdr = self._nox_edf_header
                if hdr is None:
                    raise RuntimeError("Nox EDF header not available")
                raw = read_nox_edf_channel(hdr.path, hdr, sig_idx, start_samp, n_samp)
                if as_float:
                    return raw.astype(np.float32), fs
                return raw, fs

            import pyedflib

            f = pyedflib.EdfReader(str(edf_path))
            sig = f.readSignal(sig_idx, start_samp, n_samp)
            f.close()
            if as_float:
                return np.asarray(sig, dtype=np.float32), fs
            return np.asarray(sig, dtype=np.int16), fs
        except Exception as e:
            print(f"[EDF] _get_data_edf error sig_idx={sig_idx}: {e}")
            return np.array([], dtype=np.float32 if as_float else np.int16), fs

    def get_events(self, include_xls: bool = True, custom_xls_path: Optional[str] = None) -> List[Dict]:
        """
        嘗試從 DeviceEvents.nef (SQLite) 取得事件。
        若無或解析失敗則回傳空 list。
        注意：此 DB 主要記錄裝置通知，完整臨床事件（AHI、階段等）通常在 event_*.xls / *.xlsx / docx 中。
        現在無 LIMIT，顯示所有資料（20個就20個，200個就200個）。
        include_xls=True 時會自動嘗試合併 event_*.xls / *.xlsx 裡的分析事件作為預設資料來源。
        若提供 custom_xls_path，則使用指定 XLS/XLSX 替換預設的 event 檔（用於 UI「替換XLS」功能）。
        效能優化：第一次呼叫後快取結果，避免重複解析大型 Excel 檔。
        無論來源是 .xls 或 .xlsx，解析成功後一律設定 _source='xls'，後續嚴格 exact 匹配。
        """
        cache_key = (include_xls, custom_xls_path)
        if hasattr(self, '_events_cache') and cache_key in getattr(self, '_events_cache', {}):
            return list(self._events_cache[cache_key])
        markers: List[Dict] = []
        nef = self.raw_dir / "DeviceEvents.nef"
        if nef.exists():
            try:
                conn = sqlite3.connect(str(nef))
                cur = conn.cursor()
                cur.execute(
                    "SELECT name FROM sqlite_master WHERE type='table' AND name LIKE '%marker%'"
                )
                for (tbl,) in cur.fetchall():
                    cur.execute(f"SELECT * FROM [{tbl}] LIMIT 1")
                    cols = [d[0] for d in cur.description]
                    if 'starts_at' not in cols and 'starts' not in cols:
                        continue
                    cur.execute(f"SELECT * FROM [{tbl}]")
                    for row in cur.fetchall():
                        markers.append(dict(zip(cols, row)))
                conn.close()
            except Exception as exc:
                print(f"[NEF] DeviceEvents 解析失敗: {exc}")
        if include_xls:
            markers.extend(self._parse_xls_events(xls_path=custom_xls_path))
        if not hasattr(self, '_events_cache'):
            self._events_cache = {}
        self._events_cache[cache_key] = markers
        return list(markers)

    def get_xls_events(self, xls_path: Optional[str] = None) -> List[Dict]:
        """公開 API：取得 event_*.xls / *.xlsx 裡的分析事件（Single Snore, LM, 體位等）。
        作為預設加入 events 資料，或由 UI 提供手動選擇匯入。
        若提供 xls_path，則從指定檔案解析（用於「替換XLS」功能）。
        解析成功後一律設定 _source='xls'（嚴格 exact 匹配）。
        """
        return self._parse_xls_events(xls_path=xls_path)

    def _parse_xls_events(self, xls_path: Optional[str] = None) -> list:
        """解析 event_*.xls / *.xlsx（Nox 分析軟體輸出）。"""
        if xls_path:
            chosen = [Path(xls_path)] if Path(xls_path).exists() else []
        else:
            chosen = find_event_xls_paths(self.patient_dir, self.raw_dir)
        if not chosen:
            return []
        rec_start = self.start_datetime
        if rec_start is None:
            rec_start = self._guess_start_time()
        return parse_event_xls_file(chosen[0], rec_start, fallback_start=rec_start)

    def iter_epochs(
        self,
        channels: Optional[List[str]] = None,
        epoch_sec: float = 30.0,
        step_sec: Optional[float] = None,
    ):
        """
        產生器：逐個 epoch  yield (epoch_idx, t_start, data_dict)
        data_dict: {ch_name: np.ndarray, ...}
        適合睡眠分期或特徵提取等批次處理，不會一次吃掉全部記憶體。
        """
        if step_sec is None:
            step_sec = epoch_sec
        if channels is None:
            channels = [c["name"] for c in self.channels if c["fs"] > 1][:8]  # 預設取較高頻且少量

        total = self.duration_sec
        t = 0.0
        idx = 0
        while t < total:
            data_dict = {}
            for ch in channels:
                try:
                    arr, _ = self.get_data(ch, start_sec=t, duration=epoch_sec)
                    data_dict[ch] = arr
                except Exception:
                    data_dict[ch] = np.array([])
            yield idx, t, data_dict
            t += step_sec
            idx += 1

    def __repr__(self):
        return f"<NoxRecording {self.name} duration={self.duration_sec/3600:.2f}h channels={len(self.ch_index)}>"


class NoxEdfRecording:
    """直接以 Nox 專有 EDF 檔案為資料來源的讀取器（無需 NDF / SETUP.INI）。"""

    def __init__(self, edf_path: PathLike):
        self.edf_path = Path(edf_path).resolve()
        if not self.edf_path.exists():
            raise FileNotFoundError(f"找不到 EDF 檔案：{self.edf_path}")

        self._prepare_standard_or_nox()
        self.name = self.edf_path.stem
        self.patient_dir = self.edf_path.parent
        self.raw_dir = None
        self.setup_meta: Dict[str, Dict] = {}
        self.device_info: Dict = {}
        self.ch_index: Dict[str, Dict] = {}
        self._build_channel_index()

        self.duration_sec = compute_overall_duration(self.ch_index)
        if self._hdr:
            self.start_datetime = self._hdr.start_datetime
            if self.duration_sec < _MIN_RELIABLE_NDF_DURATION_SEC:
                self.duration_sec = self._hdr.duration_sec
        else:
            edf_dur, edf_start = _read_standard_edf_meta(self.edf_path)
            if self.duration_sec < _MIN_RELIABLE_NDF_DURATION_SEC and edf_dur > 0:
                self.duration_sec = edf_dur
            self.start_datetime = edf_start
        # 單檔 EDF：所有通道共享同一起始與時間軸，無分批啟動偏移
        self.total_span_sec = self.duration_sec
        print(
            f"[NoxEdfRecording] Init {self.name}: channels={len(self.channels)}, "
            f"dur~{self.duration_sec:.1f}s, nox_format={self._hdr is not None}"
        )

    def channel_offset_sec(self, name: str) -> float:
        """單檔 EDF 所有通道共享 start_datetime，偏移恆為 0（與 NoxRecording 介面一致）。"""
        return 0.0

    def _prepare_standard_or_nox(self) -> None:
        self._hdr: Optional[NoxEdfHeader] = None
        self._edf_uses_pyedflib = False
        ok, msg = try_pyedflib_open(self.edf_path)
        if ok:
            self._edf_uses_pyedflib = True
            print(f"[NoxEdfRecording] Using standard pyedflib reader: {msg}")
            return
        print(f"[NoxEdfRecording] pyedflib failed ({msg}), using Nox parser")
        self._hdr = parse_nox_edf(self.edf_path)

    def _build_channel_index(self) -> None:
        if self._hdr:
            for i, label in enumerate(self._hdr.labels):
                if not label:
                    continue
                fs = self._hdr.fs_for_channel(i)
                if fs <= 0:
                    continue
                _, unit = _guess_fs_unit_from_label(label)
                ns = self._hdr.total_samples_for_channel(i)
                entry = {
                    "path": str(self.edf_path),
                    "stem": label,
                    "fs": fs,
                    "unit": unit,
                    "desc": label,
                    "internal_key": None,
                    "samples": ns,
                    "duration_sec": ns / fs,
                    "header_offset": 0,
                    "inferred_fs": None,
                    "source": "edf",
                    "edf_signal_index": i,
                    "nox_edf": True,
                }
                self.ch_index[label] = entry
                nk = _normalize(label)
                if nk and nk not in self.ch_index:
                    self.ch_index[nk] = entry
            return

        if self._edf_uses_pyedflib:
            import pyedflib

            f = pyedflib.EdfReader(str(self.edf_path))
            for i in range(f.signals_in_file):
                label = (f.getLabel(i) or "").strip()
                if not label:
                    continue
                fs = float(f.getSampleFrequency(i) or 0)
                if fs <= 0:
                    continue
                unit = (f.getPhysicalDimension(i) or "").strip()
                ns = int(f.getNSamples()[i]) if i < len(f.getNSamples()) else 0
                entry = {
                    "path": str(self.edf_path),
                    "stem": label,
                    "fs": fs,
                    "unit": unit,
                    "desc": label,
                    "internal_key": None,
                    "samples": ns,
                    "duration_sec": ns / fs if fs else 0.0,
                    "header_offset": 0,
                    "inferred_fs": None,
                    "source": "edf",
                    "edf_signal_index": i,
                }
                self.ch_index[label] = entry
                nk = _normalize(label)
                if nk and nk not in self.ch_index:
                    self.ch_index[nk] = entry
            f.close()

    @property
    def channels(self) -> List[Dict]:
        seen = set()
        out = []
        for k, e in self.ch_index.items():
            stem = e["stem"]
            if stem in seen:
                continue
            seen.add(stem)
            out.append(
                {
                    "name": stem,
                    "fs": e["fs"],
                    "unit": e["unit"],
                    "desc": e["desc"],
                    "samples": e["samples"],
                    "duration_sec": e["duration_sec"],
                    "path": str(e["path"]),
                    "header_offset": 0,
                    "inferred_fs": None,
                    "source_label": format_channel_source(e),
                }
            )
        out.sort(key=lambda x: (-x["fs"], x["name"]))
        return out

    def list_channels(self, include_imp: bool = False) -> List[str]:
        names = []
        for c in self.channels:
            nm = c["name"]
            if not include_imp and "imp" in _normalize(nm):
                continue
            names.append(nm)
        return names

    def get_channel_info(self, name: str) -> Optional[Dict]:
        nk = _normalize(name)
        if nk in self.ch_index:
            return self.ch_index[nk].copy()
        for k, e in self.ch_index.items():
            if _normalize(k) == nk or nk in _normalize(k):
                return e.copy()
        return None

    def get_data(
        self,
        name: str,
        start_sec: float = 0.0,
        duration: Optional[float] = None,
        as_float: bool = True,
        *,
        strip_header: bool = True,
    ) -> Tuple[np.ndarray, float]:
        info = self.get_channel_info(name)
        if info is None:
            raise KeyError(f"找不到通道：{name}")
        sig_idx = info.get("edf_signal_index")
        fs = info.get("fs") or 1.0
        total = int(info.get("samples", 0))
        start_samp = max(0, int(round(start_sec * fs)))
        if duration is None or duration <= 0:
            n_samp = total - start_samp
        else:
            n_samp = int(round(duration * fs))
        n_samp = min(max(0, n_samp), total - start_samp)
        if n_samp <= 0:
            return np.array([], dtype=np.float32 if as_float else np.int16), fs

        if self._hdr:
            raw = read_nox_edf_channel(self.edf_path, self._hdr, sig_idx, start_samp, n_samp)
            return (raw.astype(np.float32) if as_float else raw), fs

        import pyedflib

        f = pyedflib.EdfReader(str(self.edf_path))
        sig = f.readSignal(sig_idx, start_samp, n_samp)
        f.close()
        if as_float:
            return np.asarray(sig, dtype=np.float32), fs
        return np.asarray(sig, dtype=np.int16), fs

    def is_position_channel(self, name: str) -> bool:
        info = self.get_channel_info(name)
        if not info:
            return False
        nm = _normalize(name)
        stem = str(info.get("stem", "")).lower()
        desc = str(info.get("desc", "")).lower()
        return stem == "position" or "position" in nm or "posangle" in nm or "posangle" in desc

    def get_events(self, include_xls: bool = True, custom_xls_path: Optional[str] = None) -> List[Dict]:
        """純 EDF 模式仍嘗試載入同資料夾的 event_*.xls / word 報告所在層級。"""
        cache_key = (include_xls, custom_xls_path)
        if hasattr(self, '_events_cache') and cache_key in getattr(self, '_events_cache', {}):
            return list(self._events_cache[cache_key])
        markers: List[Dict] = []
        if include_xls:
            markers.extend(self.get_xls_events(xls_path=custom_xls_path))
        if not hasattr(self, '_events_cache'):
            self._events_cache = {}
        self._events_cache[cache_key] = markers
        return list(markers)

    def get_xls_events(self, xls_path: Optional[str] = None) -> List[Dict]:
        if xls_path:
            chosen = [Path(xls_path)] if Path(xls_path).exists() else []
        else:
            chosen = find_event_xls_paths(self.patient_dir, None)
        if not chosen:
            return []
        return parse_event_xls_file(
            chosen[0],
            self.start_datetime,
            fallback_start=self.start_datetime,
        )

    def get_patient_info(self, mask_pii: bool = False) -> Dict:
        """優先讀取同病患資料夾的 word_*.docx；否則退回 EDF header 技術資訊。"""
        raw_dir = find_raw_data_dir(self.patient_dir)
        if raw_dir:
            try:
                rec = NoxRecording(self.patient_dir)
                return rec.get_patient_info(mask_pii=mask_pii)
            except Exception as e:
                print(f"[NoxEdfRecording] docx/NDF 臨床資訊載入失敗，使用 EDF header: {e}")
        demo: Dict = {"note": "僅 EDF 檔案，無睡眠報告"}
        if self._hdr:
            raw = (self._hdr.patient or "").strip()
            parts = raw.split()
            if parts:
                demo["mrn"] = parts[0]
            if len(parts) >= 2 and parts[1] in ("M", "F", "Male", "Female"):
                g = parts[1]
                demo["gender"] = "Male" if g.startswith("M") else "Female"
            if self._hdr.start_datetime:
                dt = self._hdr.start_datetime
                demo["study_date"] = f"{dt.year}/{dt.month}/{dt.day}"
        if not demo.get("mrn"):
            demo["mrn"] = self.name
        if mask_pii and demo.get("mrn"):
            demo["mrn"] = str(demo["mrn"])[:8] + "***"
        return {"demographics": demo, "clinical": {}, "source": str(self.edf_path)}

    def __repr__(self):
        return f"<NoxEdfRecording {self.name} duration={self.duration_sec/3600:.2f}h channels={len(self.channels)}>"


class NoxStudy:
    """
    管理多個病患資料夾（例如 input/ 下的兩個資料夾）。
    提供統一介面批次處理。
    """

    def __init__(self, root: PathLike = "input"):
        self.root = Path(root).resolve()
        self.recordings: Dict[str, Union[NoxRecording, NoxEdfRecording]] = {}
        self._discover()

    @classmethod
    def from_edf(cls, edf_path: PathLike) -> "NoxStudy":
        """從 EDF 建立 study。若 EDF 位於含 NDF 的病患資料夾，優先完整 NoxRecording（含 Excel/報告）。"""
        edf_path = Path(edf_path).resolve()
        patient_dir = edf_path.parent
        study = cls.__new__(cls)
        study.root = patient_dir
        study.recordings = {}
        if find_raw_data_dir(patient_dir):
            try:
                rec = NoxRecording(patient_dir)
                study.recordings[rec.name] = rec
                print(f"[NoxStudy.from_edf] 偵測到 NDF，改以完整資料夾模式載入：{rec.name}")
                return study
            except Exception as e:
                print(f"[NoxStudy.from_edf] NDF 載入失敗，退回純 EDF 模式: {e}")
        rec = NoxEdfRecording(edf_path)
        study.recordings[rec.name] = rec
        return study

    def _discover(self):
        # 尋找所有像 D18_... 或 D20_... 的病患資料夾
        # 同時支援使用者直接選擇單一 Dxx_ 資料夾作為 root（此時 root 本身即為 patient）
        # 重要：跳過名稱含 "raw data" / "raw_data" 的子資料夾（它們是 raw 層，EDF 通常與它同級而非在裡面）
        candidates = []
        if self.root.is_dir():
            candidates.append(self.root)
        for d in self.root.glob("*"):
            if d.is_dir():
                candidates.append(d)
        for d in candidates:
            name_l = d.name.lower()
            if "raw data" in name_l or "raw_data" in name_l:
                continue  # 絕不把 raw data 資料夾本身當成 patient_dir，否則會在裡面找 EDF 而找不到同級的 EDF
            if re.match(r"D\d+_", d.name) or "PSG" in d.name:
                key = d.name
                if key in self.recordings:
                    continue
                try:
                    rec = NoxRecording(d)
                    self.recordings[key] = rec
                except Exception as e:
                    edf_cands = sorted(d.glob("*.edf"), key=lambda p: -p.stat().st_size)
                    if edf_cands:
                        try:
                            rec = NoxEdfRecording(edf_cands[0])
                            self.recordings[key] = rec
                            print(f"[NoxStudy] {d.name}: NDF 不可用，改以 EDF 載入 ({edf_cands[0].name})")
                            continue
                        except Exception as ee:
                            print(f"[NoxStudy] 跳過 {d.name}: NDF={e}; EDF={ee}")
                    else:
                        print(f"[NoxStudy] 跳過 {d.name}: {e}")

    @property
    def patients(self) -> List[str]:
        return sorted(self.recordings.keys())

    def get(self, key: str) -> NoxRecording:
        if key in self.recordings:
            return self.recordings[key]
        # 模糊比對
        for k, rec in self.recordings.items():
            if key in k or key == rec.name:
                return rec
        raise KeyError(f"找不到病患：{key}，可用：{self.patients}")

    def __len__(self):
        return len(self.recordings)

    def __repr__(self):
        return f"<NoxStudy patients={self.patients}>"


# 方便直接使用
def load_study(root: PathLike = "input") -> NoxStudy:
    return NoxStudy(root)


# ===================== 標準 EDF 匯出功能 =====================
def export_to_standard_edf(
    recording: NoxRecording,
    channel_names: List[str],
    out_path: PathLike,
    *,
    max_duration_sec: Optional[float] = None,
    physical_range: Optional[Tuple[float, float]] = None,
) -> Path:
    """
    將指定通道從 raw .ndf 匯出為**完全標準相容**的 EDF(+) 檔案。
    這是解決原廠 .edf 不合規（Physical Maximum 錯誤導致無法被 pyedflib / MNE 開啟）的核心修復方案。

    優點：
    - 產生檔案可直接被 EDFBrowser、MNE-Python、臨床軟體開啟。
    - 只讀取需要的通道 + 支援截斷時間，記憶體安全。
    - 保留原始取樣率與單位描述（來自 SETUP.INI）。

    注意：
    - 因為原始 .ndf 儲存的是儀器內部整數值，匯出時我們以觀測到的 min/max 當作 physical range。
    - 使用 stripped data 計算 pmin/pmax 更正確（避免 junk 污染）；所有 .ndf（含 c3/spo2 等，D18/D20 junk prefix 普遍）均受益（呼叫時預設 strip_header=True）。
    - 如需精確 uV / cmH2O 校正，請提供 physical_range 或後續在分析端 scale。
    """
    import pyedflib

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    # 過濾有效通道
    valid = []
    signals_info = []
    for nm in channel_names:
        info = recording.get_channel_info(nm)
        if info and info.get("fs", 0) > 0:
            valid.append(nm)
            signals_info.append(info)

    if not valid:
        raise ValueError("沒有可匯出的有效通道")

    # 決定總時長
    dur = recording.duration_sec
    if max_duration_sec and max_duration_sec > 0:
        dur = min(dur, max_duration_sec)

    # 建立 writer（使用較大 buffer 避免頻繁 flush）
    n_channels = len(valid)
    # 計算 header 需要的 bytes（固定）
    writer = pyedflib.EdfWriter(
        str(out_path), n_channels=n_channels, file_type=pyedflib.FILETYPE_EDFPLUS
    )

    # 設定全域 header
    start_dt = recording.start_datetime or datetime(2026, 5, 11, 22, 2, 55)
    writer.setStartdatetime(start_dt)
    writer.setPatientCode(recording.name[:20])
    writer.setPatientName(recording.name)
    # pyedflib 對 equipment / transducer 有 ASCII + 無空白限制，需清理
    writer.setEquipment("NoxA1_raw_reconstructed")

    # 為每個通道設定 header 與之後才寫入資料
    channel_headers = []
    for i, (nm, info) in enumerate(zip(valid, signals_info)):
        fs = float(info["fs"])
        unit = info.get("unit") or "uV"
        # 先讀一次全通道或前一段來決定 physical min/max（避免一次吃完整個大檔）
        # 為了簡單與正確，我們讀取整個通道（但 memmap 仍很省記憶體）
        sig, _ = recording.get_data(nm, start_sec=0, duration=dur, as_float=True, strip_header=True)
        if len(sig) == 0:
            sig = np.zeros(1, dtype=np.float32)

        pmin = float(np.min(sig)) if physical_range is None else physical_range[0]
        pmax = float(np.max(sig)) if physical_range is None else physical_range[1]
        if pmin == pmax:
            pmax = pmin + 1.0

        ch_header = {
            "label": nm[:16],
            "dimension": unit[:8],
            "sample_frequency": fs,
            "physical_min": pmin,
            "physical_max": pmax,
            "digital_min": -32768,
            "digital_max": 32767,
            "prefilter": "",
            "transducer": (info.get("desc", "") or nm)[:80],
        }
        writer.setSignalHeader(i, ch_header)
        # 確保為 float64，pyedflib write 要求
        channel_headers.append((i, nm, sig.astype("float64"), fs))

    # 正確寫入方式：使用 writeSamples 一次給完整各通道 array list
    # 順序必須與 setSignalHeader 時的順序一致
    data_list = [sig for (_, _, sig, _) in channel_headers]
    writer.writeSamples(data_list)

    writer.close()
    return out_path
